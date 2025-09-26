"""
BUTLER RAG (Retrieval Augmented Generation) System
Document processing, embedding, and intelligent retrieval for Texas County operations
"""

import os
import json
import sqlite3
import hashlib
import logging
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from enum import Enum
import asyncio
import aiofiles
import numpy as np
from pathlib import Path

# Document processing
import PyPDF2
import docx
import re

# Vector operations
from sentence_transformers import SentenceTransformer
import faiss

@dataclass
class Document:
    id: str
    title: str
    content: str
    doc_type: str
    department: str
    upload_date: datetime
    metadata: Dict[str, Any]
    embedding: Optional[np.ndarray] = None
    chunks: Optional[List[str]] = None

class DocumentType(Enum):
    COURT_DOCUMENT = "court_document"
    PROCEDURAL = "procedural"
    POLICY = "policy"
    FORM_TEMPLATE = "form_template"
    LEGAL_BRIEF = "legal_brief"
    COUNTY_ORDINANCE = "ordinance"
    MEETING_MINUTES = "minutes"
    PUBLIC_NOTICE = "notice"
    FOIA_RESPONSE = "foia"

class RAGSystem:
    """
    Advanced RAG system for Texas County document intelligence
    """

    def __init__(self, butler_core=None):
        self.butler = butler_core
        self.logger = logging.getLogger('BUTLER.RAG')

        # Initialize components
        self.db_path = "documents.db"
        self.vector_index_path = "document_vectors.index"
        self.chunk_size = 500  # Characters per chunk
        self.chunk_overlap = 100

        # Initialize embedding model (using sentence-transformers)
        self.embedder = None  # Will be initialized on first use
        self.vector_index = None

        # Document templates for generation
        self.templates = {}

        # Initialize database
        self._initialize_db()
        self._load_templates()

    def _initialize_db(self):
        """Initialize document database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute('''
            CREATE TABLE IF NOT EXISTS documents (
                id TEXT PRIMARY KEY,
                title TEXT NOT NULL,
                content TEXT NOT NULL,
                doc_type TEXT,
                department TEXT,
                upload_date REAL,
                file_path TEXT,
                metadata TEXT,
                embedding BLOB,
                chunks TEXT,
                indexed BOOLEAN DEFAULT 0
            )
        ''')

        cursor.execute('''
            CREATE TABLE IF NOT EXISTS document_chunks (
                chunk_id TEXT PRIMARY KEY,
                doc_id TEXT,
                chunk_index INTEGER,
                chunk_text TEXT,
                embedding BLOB,
                FOREIGN KEY (doc_id) REFERENCES documents(id)
            )
        ''')

        cursor.execute('''
            CREATE TABLE IF NOT EXISTS queries (
                query_id TEXT PRIMARY KEY,
                query_text TEXT,
                response TEXT,
                documents_used TEXT,
                timestamp REAL,
                user_id TEXT
            )
        ''')

        conn.commit()
        conn.close()

    def _load_templates(self):
        """Load document generation templates"""
        self.templates = {
            "court_summons": {
                "title": "Court Summons",
                "template": """
COURT SUMMONS

THE STATE OF TEXAS
COUNTY OF {county}

TO: {defendant_name}
ADDRESS: {defendant_address}

YOU ARE HEREBY COMMANDED to appear before the {court_name} Court of {county} County, Texas,
at {court_address}, on {appearance_date} at {appearance_time}.

CASE NUMBER: {case_number}
PLAINTIFF: {plaintiff_name}
DEFENDANT: {defendant_name}

NATURE OF SUIT: {suit_description}

You are further notified that if you fail to appear as commanded, a default judgment may be
entered against you granting the relief demanded in the petition.

ISSUED this {issue_date}.

_____________________
CLERK OF COURT
{county} County, Texas

By: _____________________
    Deputy Clerk
"""
            },
            "public_notice": {
                "title": "Public Notice",
                "template": """
PUBLIC NOTICE

{county} COUNTY, TEXAS

NOTICE OF {notice_type}

Date: {date}
Time: {time}
Location: {location}

The {department} of {county} County hereby provides notice of {notice_subject}.

DETAILS:
{notice_details}

PUBLIC PARTICIPATION:
{participation_info}

For more information, contact:
{contact_name}
{contact_title}
{contact_phone}
{contact_email}

This notice is posted in accordance with Texas Local Government Code Chapter {code_reference}.

Posted: {post_date}
"""
            },
            "foia_response": {
                "title": "FOIA Response Letter",
                "template": """
{department}
{county} County, Texas
{department_address}

{date}

{requestor_name}
{requestor_address}

RE: Public Information Request #{request_number}

Dear {requestor_name}:

This letter is in response to your request for public information received on {request_date}.

REQUEST SUMMARY:
{request_summary}

RESPONSE:
{response_type}

{response_details}

FEES:
{fee_information}

APPEAL RIGHTS:
You may appeal this decision to the Office of the Attorney General of Texas within {appeal_days} days.

Sincerely,

{responder_name}
{responder_title}
Public Information Coordinator
"""
            }
        }

    async def process_document(self, file_path: str, doc_type: DocumentType,
                              department: str, metadata: Dict = None) -> Document:
        """
        Process and index a document for RAG
        """
        self.logger.info(f"Processing document: {file_path}")

        # Extract text based on file type
        file_ext = Path(file_path).suffix.lower()

        if file_ext == '.pdf':
            content = await self._extract_pdf_text(file_path)
        elif file_ext in ['.docx', '.doc']:
            content = await self._extract_docx_text(file_path)
        elif file_ext == '.txt':
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                content = await f.read()
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

        # Create document object
        doc_id = hashlib.md5(f"{file_path}{datetime.now()}".encode()).hexdigest()

        document = Document(
            id=doc_id,
            title=Path(file_path).stem,
            content=content,
            doc_type=doc_type.value,
            department=department,
            upload_date=datetime.now(),
            metadata=metadata or {}
        )

        # Chunk the document
        document.chunks = self._chunk_text(content)

        # Generate embeddings
        if self.embedder is None:
            self._initialize_embedder()

        # Store in database
        await self._store_document(document)

        # Update vector index
        await self._index_document(document)

        self.logger.info(f"Document processed: {doc_id} with {len(document.chunks)} chunks")

        return document

    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
        except Exception as e:
            self.logger.error(f"Error extracting PDF text: {e}")
        return text

    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            self.logger.error(f"Error extracting DOCX text: {e}")
            return ""

    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks for better retrieval
        """
        # Clean text
        text = re.sub(r'\s+', ' ', text)

        chunks = []
        words = text.split()

        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk = ' '.join(words[i:i + self.chunk_size])
            if chunk:
                chunks.append(chunk)

        return chunks

    def _initialize_embedder(self):
        """Initialize the sentence transformer model"""
        # Using a smaller model for demo - in production use larger models
        self.embedder = SentenceTransformer('all-MiniLM-L6-v2')

        # Initialize or load FAISS index
        if os.path.exists(self.vector_index_path):
            self.vector_index = faiss.read_index(self.vector_index_path)
        else:
            # Create new index (384 dimensions for MiniLM)
            self.vector_index = faiss.IndexFlatL2(384)

    async def _store_document(self, document: Document):
        """Store document in database"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute('''
            INSERT INTO documents (id, title, content, doc_type, department,
                                 upload_date, metadata, chunks, indexed)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            document.id,
            document.title,
            document.content,
            document.doc_type,
            document.department,
            document.upload_date.timestamp(),
            json.dumps(document.metadata),
            json.dumps(document.chunks),
            False
        ))

        conn.commit()
        conn.close()

    async def _index_document(self, document: Document):
        """Add document to vector index"""
        if self.embedder is None:
            self._initialize_embedder()

        # Generate embeddings for chunks
        chunk_embeddings = self.embedder.encode(document.chunks)

        # Add to FAISS index
        self.vector_index.add(chunk_embeddings)

        # Save index
        faiss.write_index(self.vector_index, self.vector_index_path)

        # Store chunk embeddings in database
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        for i, (chunk, embedding) in enumerate(zip(document.chunks, chunk_embeddings)):
            chunk_id = f"{document.id}_chunk_{i}"
            cursor.execute('''
                INSERT INTO document_chunks (chunk_id, doc_id, chunk_index,
                                           chunk_text, embedding)
                VALUES (?, ?, ?, ?, ?)
            ''', (
                chunk_id,
                document.id,
                i,
                chunk,
                embedding.tobytes()
            ))

        # Mark document as indexed
        cursor.execute('''
            UPDATE documents SET indexed = 1 WHERE id = ?
        ''', (document.id,))

        conn.commit()
        conn.close()

    async def query_documents(self, query: str, top_k: int = 5,
                             department: str = None,
                             doc_type: str = None) -> List[Dict[str, Any]]:
        """
        Query documents using semantic search
        """
        if self.embedder is None:
            self._initialize_embedder()

        # Generate query embedding
        query_embedding = self.embedder.encode([query])[0]

        # Search in FAISS index
        distances, indices = self.vector_index.search(
            query_embedding.reshape(1, -1), top_k
        )

        # Retrieve relevant chunks from database
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        results = []
        for idx, distance in zip(indices[0], distances[0]):
            cursor.execute('''
                SELECT c.chunk_text, c.doc_id, d.title, d.department, d.doc_type
                FROM document_chunks c
                JOIN documents d ON c.doc_id = d.id
                WHERE c.chunk_index = ?
            ''', (int(idx),))

            row = cursor.fetchone()
            if row:
                chunk_text, doc_id, title, dept, dtype = row

                # Apply filters if specified
                if department and dept != department:
                    continue
                if doc_type and dtype != doc_type:
                    continue

                results.append({
                    'document_id': doc_id,
                    'title': title,
                    'department': dept,
                    'doc_type': dtype,
                    'chunk': chunk_text,
                    'relevance_score': float(1 / (1 + distance))
                })

        conn.close()

        return results

    async def generate_document(self, template_name: str,
                               variables: Dict[str, str]) -> str:
        """
        Generate a document from template
        """
        if template_name not in self.templates:
            raise ValueError(f"Template '{template_name}' not found")

        template = self.templates[template_name]['template']

        # Fill in variables
        for key, value in variables.items():
            template = template.replace(f"{{{key}}}", str(value))

        # Check for missing variables
        missing = re.findall(r'\{(\w+)\}', template)
        if missing:
            self.logger.warning(f"Missing variables in template: {missing}")

        return template

    async def answer_question(self, question: str, context_docs: List[Dict],
                             use_ollama: bool = True) -> str:
        """
        Answer a question using retrieved documents and Ollama
        """
        # Build context from retrieved documents
        context = "\n\n".join([
            f"Document: {doc['title']}\n{doc['chunk']}"
            for doc in context_docs[:3]  # Use top 3 chunks
        ])

        if use_ollama and self.butler and hasattr(self.butler, 'ai_engine'):
            # Use Ollama for answer generation
            prompt = f"""Based on the following Texas County documents, answer the question.

Context Documents:
{context}

Question: {question}

Provide a detailed answer based on the documents. If the documents don't contain enough information, say so."""

            try:
                response = await self.butler.ai_engine.generate_response(
                    prompt,
                    "You are BUTLER, an AI assistant for Texas County operations. Answer based on the provided documents."
                )
                return response.content
            except Exception as e:
                self.logger.error(f"Ollama error: {e}")
                return self._fallback_answer(question, context_docs)
        else:
            return self._fallback_answer(question, context_docs)

    def _fallback_answer(self, question: str, context_docs: List[Dict]) -> str:
        """
        Fallback answer generation without LLM
        """
        if not context_docs:
            return "I couldn't find any relevant documents to answer your question."

        # Simple extraction-based answer
        answer = f"Based on the documents in our system:\n\n"

        for i, doc in enumerate(context_docs[:3], 1):
            answer += f"{i}. From '{doc['title']}' ({doc['department']}):\n"
            answer += f"   {doc['chunk'][:200]}...\n\n"

        answer += f"\nThese documents may contain the information you're looking for regarding: {question}"

        return answer

    async def get_document_stats(self) -> Dict[str, Any]:
        """
        Get statistics about indexed documents
        """
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute('SELECT COUNT(*) FROM documents')
        total_docs = cursor.fetchone()[0]

        cursor.execute('SELECT COUNT(*) FROM documents WHERE indexed = 1')
        indexed_docs = cursor.fetchone()[0]

        cursor.execute('SELECT COUNT(*) FROM document_chunks')
        total_chunks = cursor.fetchone()[0]

        cursor.execute('''
            SELECT doc_type, COUNT(*)
            FROM documents
            GROUP BY doc_type
        ''')
        doc_types = dict(cursor.fetchall())

        cursor.execute('''
            SELECT department, COUNT(*)
            FROM documents
            GROUP BY department
        ''')
        departments = dict(cursor.fetchall())

        conn.close()

        return {
            'total_documents': total_docs,
            'indexed_documents': indexed_docs,
            'total_chunks': total_chunks,
            'document_types': doc_types,
            'departments': departments,
            'index_size': os.path.getsize(self.vector_index_path) if os.path.exists(self.vector_index_path) else 0
        }

# Texas County specific document processors
class TexasCountyProcessor:
    """
    Specialized processors for Texas County documents
    """

    @staticmethod
    def process_court_filing(content: str) -> Dict[str, Any]:
        """Extract structured data from court filings"""
        extracted = {
            'case_number': None,
            'parties': [],
            'filing_date': None,
            'document_type': None,
            'judge': None
        }

        # Extract case number
        case_pattern = r'(?:Case|Cause)\s*(?:No|Number)[:\s]*([A-Z0-9\-]+)'
        match = re.search(case_pattern, content, re.IGNORECASE)
        if match:
            extracted['case_number'] = match.group(1)

        # Extract parties
        plaintiff_pattern = r'Plaintiff[:\s]*([^\n]+)'
        defendant_pattern = r'Defendant[:\s]*([^\n]+)'

        p_match = re.search(plaintiff_pattern, content, re.IGNORECASE)
        d_match = re.search(defendant_pattern, content, re.IGNORECASE)

        if p_match:
            extracted['parties'].append({'role': 'plaintiff', 'name': p_match.group(1).strip()})
        if d_match:
            extracted['parties'].append({'role': 'defendant', 'name': d_match.group(1).strip()})

        return extracted

    @staticmethod
    def process_listserv_email(content: str) -> Dict[str, Any]:
        """Process Texas County ListServ emails"""
        extracted = {
            'sender_department': None,
            'priority': 'normal',
            'action_required': False,
            'deadline': None,
            'topics': []
        }

        # Check for priority markers
        if any(word in content.lower() for word in ['urgent', 'emergency', 'immediate']):
            extracted['priority'] = 'high'

        # Check for action requirements
        if any(phrase in content.lower() for phrase in ['action required', 'response needed', 'please respond']):
            extracted['action_required'] = True

        # Extract deadline
        deadline_pattern = r'(?:deadline|due|by)[:\s]*([^\n]+(?:202\d|pm|am))'
        match = re.search(deadline_pattern, content, re.IGNORECASE)
        if match:
            extracted['deadline'] = match.group(1).strip()

        return extracted